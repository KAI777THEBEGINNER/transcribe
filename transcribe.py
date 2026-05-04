"""
Local English dialogue transcription pipeline.
- Step 1: faster-whisper transcription (medium.en, int8)
- Step 2: wav2vec2 word-level alignment
- Step 3: pyannote speaker diarization
- Step 4: assign speakers + render .docx

All intermediate results are written to a hidden work dir next to the audio,
so any step can resume after a crash.
"""
import argparse
import gc
import json
import os
import shutil
import sys
import time
from datetime import datetime, timedelta
from pathlib import Path

from dotenv import load_dotenv

SCRIPT_DIR = Path(__file__).resolve().parent
load_dotenv(SCRIPT_DIR / ".env")

HF_TOKEN = os.environ.get("HF_TOKEN")
HF_CACHE_DIR = os.environ.get("HF_CACHE_DIR", str(SCRIPT_DIR / "hf_cache"))
TORCH_CACHE_DIR = os.environ.get("TORCH_CACHE_DIR", str(Path(HF_CACHE_DIR).parent / "torch_cache"))
PROXY = os.environ.get("PROXY", "")

if not HF_TOKEN:
    sys.exit("HF_TOKEN missing — put it in .env")

# Must be set BEFORE importing whisperx / huggingface_hub / torchaudio
os.environ["HF_HOME"] = HF_CACHE_DIR
os.environ["HUGGINGFACE_HUB_CACHE"] = HF_CACHE_DIR
os.environ["TRANSFORMERS_CACHE"] = HF_CACHE_DIR
os.environ["TORCH_HOME"] = TORCH_CACHE_DIR  # torchaudio wav2vec2 lives here
os.environ["HF_HUB_DISABLE_SYMLINKS"] = "1"  # exFAT compat
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
os.environ["TOKENIZERS_PARALLELISM"] = "false"
if PROXY:
    os.environ.setdefault("HTTP_PROXY", PROXY)
    os.environ.setdefault("HTTPS_PROXY", PROXY)
    os.environ.setdefault("http_proxy", PROXY)
    os.environ.setdefault("https_proxy", PROXY)

Path(HF_CACHE_DIR).mkdir(parents=True, exist_ok=True)
Path(TORCH_CACHE_DIR).mkdir(parents=True, exist_ok=True)


def fmt_ts(seconds: float) -> str:
    return str(timedelta(seconds=int(seconds)))


def fmt_speaker(raw: str) -> str:
    if not raw:
        return "??"
    return raw.replace("SPEAKER_", "").lstrip("0") or "0"


def step1_transcribe(args, audio_path: Path, seg_path: Path):
    print(f"\n[1/4] faster-whisper {args.model} transcribing... (this is the slow one)")
    t0 = time.time()
    import whisperx

    model = whisperx.load_model(
        args.model,
        device="cpu",
        compute_type=args.compute_type,
        language=args.lang,
        download_root=HF_CACHE_DIR,
    )
    audio = whisperx.load_audio(str(audio_path))
    result = model.transcribe(audio, batch_size=args.batch_size, language=args.lang)

    seg_path.write_text(json.dumps(result, ensure_ascii=False, indent=2, default=str))
    del model
    gc.collect()
    print(f"      done in {time.time() - t0:.0f}s -> {seg_path.name}")
    return result, audio


def step2_align(args, audio, result, align_path: Path):
    print(f"\n[2/4] word-level alignment...")
    t0 = time.time()
    import whisperx

    model_a, metadata = whisperx.load_align_model(
        language_code=args.lang, device="cpu"
    )
    aligned = whisperx.align(
        result["segments"],
        model_a,
        metadata,
        audio,
        device="cpu",
        return_char_alignments=False,
    )
    align_path.write_text(json.dumps(aligned, ensure_ascii=False, indent=2, default=str))
    del model_a
    gc.collect()
    print(f"      done in {time.time() - t0:.0f}s -> {align_path.name}")
    return aligned


def step3_diarize(args, audio_path: Path, audio, diar_path: Path):
    print(f"\n[3/4] pyannote speaker diarization...")
    t0 = time.time()
    try:
        from whisperx import DiarizationPipeline
    except ImportError:
        from whisperx.diarize import DiarizationPipeline

    diar_model = DiarizationPipeline(use_auth_token=HF_TOKEN, device="cpu")
    diarize_segments = diar_model(
        str(audio_path) if audio is None else audio,
        min_speakers=args.min_speakers,
        max_speakers=args.max_speakers,
    )

    if hasattr(diarize_segments, "to_json"):
        diar_path.write_text(diarize_segments.to_json(orient="records", indent=2))
    else:
        diar_path.write_text(json.dumps(diarize_segments, ensure_ascii=False, indent=2, default=str))

    del diar_model
    gc.collect()
    print(f"      done in {time.time() - t0:.0f}s -> {diar_path.name}")
    return diarize_segments


def step4_assign(aligned, diarize_segments, final_path: Path):
    print(f"\n[4/4] merging text and speaker labels...")
    t0 = time.time()
    import whisperx

    final = whisperx.assign_word_speakers(diarize_segments, aligned)
    final_path.write_text(json.dumps(final, ensure_ascii=False, indent=2, default=str))
    print(f"      done in {time.time() - t0:.0f}s -> {final_path.name}")
    return final


def render_docx(audio_path: Path, final: dict) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor

    out_path = audio_path.with_suffix(".docx")
    doc = Document()

    title = audio_path.stem.replace("_", " ")
    doc.add_heading(title, level=1)

    meta = doc.add_paragraph()
    run = meta.add_run(f"Transcribed {datetime.now():%Y-%m-%d %H:%M}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    last_speaker = None
    current_para = None

    for seg in final.get("segments", []):
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        speaker = seg.get("speaker") or "??"
        start = seg.get("start", 0)

        if speaker != last_speaker:
            head = doc.add_paragraph()
            head_run = head.add_run(f"Speaker {fmt_speaker(speaker)}  ·  {fmt_ts(start)}")
            head_run.bold = True
            head_run.font.size = Pt(11)
            current_para = doc.add_paragraph()
            last_speaker = speaker

        current_para.add_run(text + " ")

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="Local two-speaker English transcription")
    ap.add_argument("audio", help="path to audio file (m4a/mp3/wav...)")
    ap.add_argument("--model", default="medium.en",
                    help="whisper model: tiny.en/base.en/small.en/medium.en/large-v3 (default: medium.en)")
    ap.add_argument("--lang", default="en")
    ap.add_argument("--batch-size", type=int, default=4,
                    help="lower = less memory, slower (default: 4 for 8GB Mac)")
    ap.add_argument("--compute-type", default="int8", choices=["int8", "float32"],
                    help="int8 = lighter, float32 = fp baseline")
    ap.add_argument("--min-speakers", type=int, default=2)
    ap.add_argument("--max-speakers", type=int, default=2)
    ap.add_argument("--keep-temp", action="store_true",
                    help="keep intermediate .json files for inspection")
    ap.add_argument("--from-step", type=int, default=1, choices=[1, 2, 3, 4],
                    help="resume from a step (use intermediate .json files)")
    args = ap.parse_args()

    audio_path = Path(args.audio).expanduser().resolve()
    if not audio_path.exists():
        sys.exit(f"audio not found: {audio_path}")

    work_dir = audio_path.parent / f".transcribe_{audio_path.stem}"
    work_dir.mkdir(exist_ok=True)
    seg_path = work_dir / "1_segments.json"
    align_path = work_dir / "2_aligned.json"
    diar_path = work_dir / "3_diarize.json"
    final_path = work_dir / "4_final.json"

    print(f"audio       : {audio_path}")
    print(f"work_dir    : {work_dir}")
    print(f"hf_cache    : {HF_CACHE_DIR}")
    print(f"model       : {args.model}  compute={args.compute_type}  batch={args.batch_size}")
    print(f"speakers    : min={args.min_speakers} max={args.max_speakers}")

    audio = None
    result = None
    aligned = None
    diarize_segments = None

    if args.from_step <= 1:
        result, audio = step1_transcribe(args, audio_path, seg_path)
    else:
        print(f"\n[1/4] resume: loading {seg_path.name}")
        result = json.loads(seg_path.read_text())

    if args.from_step <= 2:
        if audio is None:
            import whisperx
            audio = whisperx.load_audio(str(audio_path))
        aligned = step2_align(args, audio, result, align_path)
    else:
        print(f"\n[2/4] resume: loading {align_path.name}")
        aligned = json.loads(align_path.read_text())

    if args.from_step <= 3:
        if audio is None:
            import whisperx
            audio = whisperx.load_audio(str(audio_path))
        diarize_segments = step3_diarize(args, audio_path, audio, diar_path)
    else:
        print(f"\n[3/4] resume: loading {diar_path.name}")
        import pandas as pd
        diarize_segments = pd.read_json(diar_path, orient="records")

    final = step4_assign(aligned, diarize_segments, final_path)
    out_path = render_docx(audio_path, final)

    print(f"\nDONE -> {out_path}")
    print(f"size: {out_path.stat().st_size // 1024} KB")

    if not args.keep_temp:
        shutil.rmtree(work_dir, ignore_errors=True)
        print(f"cleaned up: {work_dir.name}")
    else:
        print(f"intermediate files kept in: {work_dir}")


if __name__ == "__main__":
    main()
